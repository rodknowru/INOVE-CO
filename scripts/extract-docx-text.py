#!/usr/bin/env python3
"""Текст из .docx с сохранением списков (упрощённо).

Правила:
- Первая непустая строка: как есть (для title).
- Маркированные/нумерованные пункты: префикс "• ".

Вызывается из sync-products.mjs.
"""
import sys

import docx


def is_list_paragraph(p) -> bool:
    # python-docx не даёт списка напрямую, но numPr в XML — хороший признак.
    try:
        ppr = p._p.pPr  # noqa: SLF001
        return ppr is not None and ppr.numPr is not None
    except Exception:
        return False


def paragraph_plain_text(p) -> str:
    return (p.text or "").strip()


def paragraph_text(p, *, is_title: bool) -> str:
    if is_title:
        return (p.text or "").strip()

    t = paragraph_plain_text(p)
    if not t:
        return ""
    if is_list_paragraph(p):
        return f"• {t}"
    return t


def main() -> None:
    path = sys.argv[1]
    d = docx.Document(path)
    parts: list[str] = []

    title_done = False
    for p in d.paragraphs:
        t = paragraph_text(p, is_title=not title_done)
        if t:
            parts.append(t)
            title_done = True

    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = paragraph_text(p, is_title=False)
                    if t:
                        parts.append(t)
    print("\n".join(parts), end="")


if __name__ == "__main__":
    main()
